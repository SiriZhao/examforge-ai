from contextlib import redirect_stderr, redirect_stdout
import csv
from html import escape
from io import StringIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.schemas.review import ExportFormat, ReviewReport
from app.services.review_planner import sanitize_report


class ExportError(RuntimeError):
    pass


def write_output(output_dir: Path, file_id: str, content: str, suffix: str = ".md") -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{file_id}{suffix}"
    path.write_text(content, encoding="utf-8")
    return path


def export_review_report(
    report: ReviewReport,
    markdown: str,
    output_dir: Path,
    basename: str,
    export_format: ExportFormat,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    if export_format == "md":
        return export_markdown(markdown, output_dir, basename)
    if export_format == "docx":
        return export_docx(report, output_dir, basename)
    if export_format == "pdf":
        return export_pdf(markdown, output_dir, basename)
    raise ExportError(f"不支持的导出格式：{export_format}。")


def export_markdown(markdown: str, output_dir: Path, basename: str) -> Path:
    path = output_dir / f"{basename}.md"
    path.write_text(markdown, encoding="utf-8")
    return path


def export_docx(report: ReviewReport, output_dir: Path, basename: str) -> Path:
    report = sanitize_report(report)
    document = Document()
    set_document_chinese_font(document)

    document.add_heading(report.title, level=1)
    document.add_paragraph(f"生成时间：{report.generated_at}")
    document.add_heading("总览", level=2)
    document.add_paragraph(report.summary)

    document.add_heading("章节复习建议", level=2)
    for chapter in report.chapters:
        document.add_heading(chapter.chapter, level=3)
        table = document.add_table(rows=7, cols=2)
        table.style = "Table Grid"
        rows = [
            ("重要度", str(chapter.importance)),
            ("出现频次", str(chapter.frequency)),
            ("关键词", "、".join(chapter.keywords) or "暂无"),
            ("公式", "、".join(chapter.formulas) or "暂无"),
            ("题型", "、".join(chapter.question_types) or "暂无"),
            ("复习建议", chapter.review_advice),
            ("示例题", "\n".join(chapter.examples) or "暂无"),
        ]
        for index, (name, value) in enumerate(rows):
            table.cell(index, 0).text = name
            table.cell(index, 1).text = value

    add_bullets(document, "高频考点", report.high_frequency_points)
    add_bullets(document, "冲刺清单", report.sprint_checklist)
    add_bullets(document, "低优先级内容", report.low_priority)
    add_bullets(document, "材料不足提示", report.insufficient_materials or ["暂无"])

    add_exam_intelligence_docx_sections(document, report)
    set_document_chinese_font(document)
    path = output_dir / f"{basename}.docx"
    document.save(path)
    return path


def export_pdf(markdown: str, output_dir: Path, basename: str) -> Path:
    html = markdown_to_html(markdown)
    path = output_dir / f"{basename}.pdf"
    try:
        with redirect_stdout(StringIO()), redirect_stderr(StringIO()):
            from weasyprint import HTML
    except Exception as exc:
        return export_pdf_with_reportlab(markdown, path, exc)

    try:
        HTML(string=html).write_pdf(path)
    except Exception as exc:
        return export_pdf_with_reportlab(markdown, path, exc)
    return path


def export_anki_csv(report: ReviewReport, output_dir: Path, basename: str) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{basename}-anki.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Front", "Back", "Tags"])
        for card in report.anki_cards:
            writer.writerow([card.front, card.back, card.tags])
    return path


def add_exam_intelligence_docx_sections(document: Document, report: ReviewReport) -> None:
    document.add_heading("往年题高频考点分析", level=2)
    document.add_paragraph(report.past_exam_analysis.summary or "暂无往年题分析。")
    for topic in report.past_exam_analysis.high_frequency_topics:
        document.add_paragraph(
            f"{topic.topic}（{topic.chapter}）- 频次 {topic.frequency}；"
            f"题型：{'、'.join(topic.question_types) or '未知'}",
            style="List Bullet",
        )

    document.add_heading("推荐复习顺序", level=2)
    for item in report.review_order:
        document.add_paragraph(
            f"{item.chapter}（{item.importance}/100）：{item.reason}",
            style="List Number",
        )

    document.add_heading("考前冲刺计划", level=2)
    for plan in report.sprint_plans:
        document.add_heading(plan.title, level=3)
        for item in plan.schedule:
            document.add_paragraph(item, style="List Bullet")

    document.add_heading("模拟卷", level=2)
    for index, question in enumerate(report.mock_exam.questions, start=1):
        document.add_heading(f"{index}. {question.question_type}", level=3)
        document.add_paragraph(question.question)
        document.add_paragraph(f"参考答案：{question.answer}")

    document.add_heading("Anki 卡片预览", level=2)
    for card in report.anki_cards[:12]:
        document.add_paragraph(f"正面：{card.front}\n背面：{card.back}\n标签：{card.tags}")


def markdown_to_html(markdown: str) -> str:
    try:
        import markdown as markdown_lib

        body = markdown_lib.markdown(markdown, extensions=["tables", "extra"])
    except Exception:
        body = f"<pre>{escape(markdown)}</pre>"

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <style>
    body {{
      font-family: "Microsoft YaHei", "SimSun", "Noto Sans CJK SC", sans-serif;
      line-height: 1.65;
      color: #111827;
      margin: 32px;
    }}
    h1, h2, h3 {{ color: #0f172a; }}
    table {{ border-collapse: collapse; width: 100%; margin: 12px 0 20px; }}
    th, td {{ border: 1px solid #cbd5e1; padding: 8px; vertical-align: top; }}
    th {{ background: #f1f5f9; }}
    code, pre {{ font-family: Consolas, "Microsoft YaHei", monospace; }}
  </style>
</head>
<body>{body}</body>
</html>"""


def set_document_chinese_font(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei" if style_name.startswith("Heading") else "SimSun"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)


def add_bullets(document: Document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=2)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def export_pdf_with_reportlab(markdown: str, path: Path, original_error: Exception) -> Path:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.pdfmetrics import registerFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except Exception as exc:
        raise ExportError(
            "PDF 导出失败。WeasyPrint 系统依赖缺失，且 ReportLab 备用方案不可用。"
        ) from exc

    try:
        registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        base_style = ParagraphStyle(
            "ChineseBody",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=10,
            leading=16,
        )
        heading_style = ParagraphStyle(
            "ChineseHeading",
            parent=styles["Heading2"],
            fontName="STSong-Light",
            fontSize=16,
            leading=22,
            spaceAfter=8,
        )
        story = []
        pending_table: list[list[str]] = []

        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                flush_markdown_table(story, pending_table, base_style, Paragraph, Table, TableStyle, colors)
                story.append(Spacer(1, 6))
                continue

            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [cell.strip() for cell in stripped.strip("|").split("|")]
                if not all(set(cell) <= {"-", ":", " "} for cell in cells):
                    pending_table.append(cells)
                continue

            flush_markdown_table(story, pending_table, base_style, Paragraph, Table, TableStyle, colors)
            if stripped.startswith("#"):
                text = stripped.lstrip("#").strip()
                story.append(Paragraph(escape(text), heading_style))
            elif stripped.startswith("- "):
                story.append(Paragraph(f"• {escape(stripped[2:])}", base_style))
            else:
                story.append(Paragraph(escape(stripped), base_style))

        flush_markdown_table(story, pending_table, base_style, Paragraph, Table, TableStyle, colors)
        document = SimpleDocTemplate(str(path), pagesize=A4)
        document.build(story)
    except Exception as exc:
        raise ExportError(f"使用 ReportLab 备用方案导出 PDF 时失败：{original_error}") from exc

    return path


def flush_markdown_table(
    story,
    rows: list[list[str]],
    paragraph_style,
    paragraph_cls,
    table_cls,
    table_style_cls,
    colors_module,
) -> None:
    if not rows:
        return

    table_data = [
        [paragraph_cls(escape(cell), paragraph_style) for cell in row]
        for row in rows
    ]
    table = table_cls(table_data, repeatRows=1)
    table.setStyle(
        table_style_cls(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("BACKGROUND", (0, 0), (-1, 0), colors_module.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors_module.lightgrey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)
    rows.clear()
